"""
JD Gap Analyzer - 简历与JD深度差距分析工具
基于 SpeedyJob 改进方案 MVP
作者：liuzh | vibe coding with AI
"""

import streamlit as st
import json
import re
import io
from openai import OpenAI
import pdfplumber
import docx

# ── 页面配置 ──────────────────────────────────────────────
st.set_page_config(
    page_title="JD Gap Analyzer · 简历差距诊断",
    page_icon="🔍",
    layout="wide",
    initial_sidebar_state="collapsed",
)

# ── 样式 ─────────────────────────────────────────────────
st.markdown("""
<style>
  .main { background: #f7f8fc; }
  .stTextArea textarea {
    font-size: 14px !important;
    min-height: 220px !important;
  }
  .summary-text {
    font-size: 15px;
    color: #1f2937;
    line-height: 1.7;
    margin-top: 12px;
  }
  .section-title {
    font-size: 16px;
    font-weight: 700;
    color: #111827;
  }
  .card {
    background: white;
    border-radius: 12px;
    padding: 20px 24px;
    margin: 10px 0;
    box-shadow: 0 2px 8px rgba(0,0,0,0.07);
  }
  .hit   { border-left: 4px solid #22c55e; }
  .gap   { border-left: 4px solid #ef4444; }
  .weak  { border-left: 4px solid #f59e0b; }
  .tag-hit  { background:#dcfce7; color:#16a34a; padding:2px 10px; border-radius:20px; font-size:12px; font-weight:600; }
  .tag-gap  { background:#fee2e2; color:#dc2626; padding:2px 10px; border-radius:20px; font-size:12px; font-weight:600; }
  .tag-weak { background:#fef9c3; color:#b45309; padding:2px 10px; border-radius:20px; font-size:12px; font-weight:600; }
  .score-circle {
    font-size: 56px;
    font-weight: 800;
    text-align: center;
    padding: 20px 0 8px;
  }
  .rewrite-box {
    background: #f0f9ff;
    border: 1px solid #bae6fd;
    border-radius: 8px;
    padding: 12px 16px;
    margin-top: 8px;
    font-size: 13px;
    color: #0c4a6e;
  }
  h1 { font-size: 26px !important; }
</style>
""", unsafe_allow_html=True)

# ── 标题 ─────────────────────────────────────────────────
st.markdown("## 🔍 JD Gap Analyzer")
st.markdown("**逐条拆解 JD 要求 · 精准定位简历短板 · 给出可直接使用的改写建议**")
st.divider()

# ── 侧边栏：API Key ───────────────────────────────────────
with st.sidebar:
    st.header("⚙️ 配置")
    api_key = st.text_input("DeepSeek API Key", type="password",
                             help="从 platform.deepseek.com 获取，本地运行不会上传")
    model_choice = st.selectbox("模型", ["deepseek-chat", "deepseek-reasoner"])
    st.caption("💡 API Key 仅在本次会话中使用，不会被存储")

# ── 文件解析函数 ──────────────────────────────────────────
def parse_pdf(file_bytes: bytes) -> str:
    text_parts = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            t = page.extract_text()
            if t:
                text_parts.append(t)
    return "\n".join(text_parts)


def parse_docx(file_bytes: bytes) -> str:
    doc = docx.Document(io.BytesIO(file_bytes))
    paragraphs = []
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text)
    # 尝试读取表格
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)
    return "\n".join(paragraphs)


def parse_file(uploaded_file) -> str | None:
    fname = uploaded_file.name.lower()
    try:
        if fname.endswith(".pdf"):
            return parse_pdf(uploaded_file.read())
        elif fname.endswith(".docx"):
            return parse_docx(uploaded_file.read())
        elif fname.endswith(".txt"):
            return uploaded_file.read().decode("utf-8", errors="replace")
        else:
            return None
    except Exception as e:
        return None


# ── 主输入区 ──────────────────────────────────────────────
col1, col2 = st.columns(2, gap="large")

# ── 简历输入区（支持文件上传 + 粘贴）─────────────────────
with col1:
    st.markdown("### 📄 简历")
    resume_file = st.file_uploader(
        "上传简历（PDF/Word/TXT）",
        type=["pdf", "docx", "txt"],
        key="resume_upload",
        help="支持 PDF、Word（.docx）、纯文本，自动提取内容"
    )
    if resume_file:
        parsed = parse_file(resume_file)
        if parsed:
            st.session_state["resume_text"] = parsed
            st.success(f"✅ 已解析：{resume_file.name}（{len(parsed)} 字）")
        else:
            st.error("文件解析失败，请尝试复制粘贴内容")
    resume_text = st.text_area(
        label="简历内容",
        height=280,
        key="resume_text",
        placeholder="上传文件后内容自动填入此处，也可直接粘贴...\n\n快捷键：Ctrl+V 粘贴",
        label_visibility="collapsed"
    )

# ── JD 输入区（支持文件上传 + 粘贴）──────────────────────
with col2:
    st.markdown("### 📋 目标 JD")
    jd_file = st.file_uploader(
        "上传 JD（PDF/Word/TXT）",
        type=["pdf", "docx", "txt"],
        key="jd_upload",
        help="支持 PDF、Word（.docx）、纯文本"
    )
    if jd_file:
        parsed = parse_file(jd_file)
        if parsed:
            st.session_state["jd_text"] = parsed
            st.success(f"✅ 已解析：{jd_file.name}（{len(parsed)} 字）")
        else:
            st.error("文件解析失败，请尝试复制粘贴内容")
    jd_text = st.text_area(
        label="JD内容",
        height=280,
        key="jd_text",
        placeholder="上传文件后内容自动填入此处，也可直接粘贴...",
        label_visibility="collapsed"
    )

# ── 示例数据 ──────────────────────────────────────────────
with st.expander("📖 点击加载示例数据（快速体验）"):
    if st.button("加载示例", use_container_width=True):
        st.session_state["resume_text"] = """工作经验
2023.06 - 至今 | 某科技公司 | Python后端工程师
• 负责用户中心、订单系统的 API 开发，使用 FastAPI + PostgreSQL
• 设计并实现了基于 Redis 的缓存层，接口响应时间降低 40%
• 参与微服务拆分，将单体应用拆解为 5 个独立服务
• 编写单元测试，覆盖率达到 85%

2022.07 - 2023.05 | 某创业公司 | 全栈实习生  
• 使用 Vue3 开发后台管理系统前端
• 协助搭建 CI/CD 流水线（GitHub Actions + Docker）

技能
编程语言：Python（熟练）、JavaScript（熟悉）
框架：FastAPI, Vue3, React（了解）
数据库：PostgreSQL, MySQL, Redis
工具：Docker, Git, Linux
英语：CET-6（阅读无障碍）"""

        st.session_state["jd_text"] = """岗位要求（AI应用开发工程师）
1. 熟悉 Python，有扎实的工程开发能力
2. 熟练使用 LangChain 或类似 LLM 框架进行 AI 应用开发
3. 有 RAG（检索增强生成）系统的设计与实现经验
4. 了解主流大语言模型 API（OpenAI、DeepSeek等）的调用方式
5. 有前后端全栈开发能力，能独立完成产品原型
6. 具备良好的英文文档阅读能力
7. 有 AI 项目上线经验者优先"""

        st.rerun()

# 加载示例数据（通过 session_state 自动同步到 text_area）

st.divider()

# ── 分析按钮 ──────────────────────────────────────────────
analyze_btn = st.button(
    "🚀 开始深度分析", 
    type="primary", 
    use_container_width=True,
    disabled=not (resume_text and jd_text and api_key)
)

if not api_key:
    st.info("👈 请在左侧侧边栏输入 DeepSeek API Key 后开始分析")

# ── AI 分析核心逻辑 ────────────────────────────────────────
SYSTEM_PROMPT = """你是一位资深的职业顾问和简历优化专家。
你的任务是对求职者的简历和目标职位JD进行深度对比分析。

分析维度：
1. 逐条拆解JD中每一项要求
2. 判断简历是否有效覆盖该要求（三个等级：hit/weak/gap）
   - hit：简历中有明确的相关内容，覆盖良好
   - weak：简历中有相关内容但不够突出或描述不够有力
   - gap：简历中完全缺失该要求相关内容
3. 对每条要求给出针对性的改写建议（如果是weak或gap，给出可直接添加到简历的具体文字建议）

输出要求：严格按照JSON格式输出，不要有任何额外文字。

JSON格式如下：
{
  "overall_score": 75,
  "summary": "总体评价（2-3句话）",
  "hit_count": 3,
  "weak_count": 2,
  "gap_count": 2,
  "items": [
    {
      "jd_requirement": "JD原文要求",
      "status": "hit|weak|gap",
      "resume_evidence": "简历中的对应内容摘要（如有）",
      "suggestion": "改写建议或加分动作（hit时可为空字符串）"
    }
  ],
  "quick_wins": ["最优先要补充的3个行动建议"],
  "strengths": ["简历相对该JD的3个优势亮点"]
}"""

def run_analysis(resume: str, jd: str, api_key: str, model: str) -> dict:
    client = OpenAI(api_key=api_key, base_url="https://api.deepseek.com")
    
    user_msg = f"""请分析以下简历与JD的匹配情况：

===== 简历内容 =====
{resume}

===== 目标JD =====
{jd}

请严格按JSON格式输出分析结果。"""

    response = client.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": SYSTEM_PROMPT},
            {"role": "user", "content": user_msg}
        ],
        temperature=0.3,
        max_tokens=3000,
    )
    
    raw = response.choices[0].message.content.strip()
    # 提取 JSON（防止模型输出多余文字）
    json_match = re.search(r'\{.*\}', raw, re.DOTALL)
    if json_match:
        return json.loads(json_match.group())
    return json.loads(raw)


def score_color(score: int) -> str:
    if score >= 80: return "#22c55e"
    if score >= 60: return "#f59e0b"
    return "#ef4444"


def _esc(text: str) -> str:
    """将文本中的换行转为<br>，转义HTML特殊字符，防止被markdown解析"""
    if not text:
        return ""
    text = str(text)
    text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    text = text.replace("\n", "<br>").replace("\r", "")
    return text


def status_tag(status: str) -> str:
    if status == "hit":
        return '<span class="tag-hit">✅ 覆盖</span>'
    elif status == "weak":
        return '<span class="tag-weak">⚠️ 较弱</span>'
    else:
        return '<span class="tag-gap">❌ 缺失</span>'


# ── 执行分析 & 渲染结果 ────────────────────────────────────
if analyze_btn:
    with st.spinner("🤖 AI 正在逐条分析 JD 要求，请稍候..."):
        try:
            result = run_analysis(resume_text, jd_text, api_key, model_choice)
        except Exception as e:
            st.error(f"分析失败：{e}")
            st.stop()

    st.success("✅ 分析完成！")
    st.divider()

    # ── 顶部总览 ──────────────────────────────────────────
    st.markdown("## 📊 总览")
    
    c1, c2, c3, c4 = st.columns(4)
    score = result.get("overall_score", 0)
    color = score_color(score)
    
    with c1:
        st.html(f"""
        <div class="card" style="text-align:center">
          <div style="font-size:13px;color:#888;margin-bottom:4px">综合匹配度</div>
          <div class="score-circle" style="color:{color}">{score}</div>
          <div style="font-size:12px;color:#888">/ 100</div>
        </div>""")
    
    with c2:
        st.html(f"""
        <div class="card" style="text-align:center">
          <div style="font-size:13px;color:#888;margin-bottom:4px">已覆盖</div>
          <div style="font-size:48px;font-weight:800;color:#22c55e;padding:20px 0 8px">
            {result.get('hit_count', 0)}
          </div>
          <div style="font-size:12px;color:#888">条 JD 要求</div>
        </div>""")
    
    with c3:
        st.html(f"""
        <div class="card" style="text-align:center">
          <div style="font-size:13px;color:#888;margin-bottom:4px">覆盖较弱</div>
          <div style="font-size:48px;font-weight:800;color:#f59e0b;padding:20px 0 8px">
            {result.get('weak_count', 0)}
          </div>
          <div style="font-size:12px;color:#888">条 JD 要求</div>
        </div>""")
    
    with c4:
        st.html(f"""
        <div class="card" style="text-align:center">
          <div style="font-size:13px;color:#888;margin-bottom:4px">完全缺失</div>
          <div style="font-size:48px;font-weight:800;color:#ef4444;padding:20px 0 8px">
            {result.get('gap_count', 0)}
          </div>
          <div style="font-size:12px;color:#888">条 JD 要求</div>
        </div>""")

    # 总结
    st.html(f"""
    <div class="card">
      <span class="section-title">📝 AI 总评</span>
      <div class="summary-text">{_esc(result.get('summary', ''))}</div>
    </div>""")

    st.divider()

    # ── 逐条分析 ──────────────────────────────────────────
    st.markdown("## 🔬 逐条 JD 要求分析")

    items = result.get("items", [])
    for i, item in enumerate(items, 1):
        status = item.get("status", "gap")
        card_class = {"hit": "hit", "weak": "weak", "gap": "gap"}.get(status, "gap")
        
        jd_req = _esc(item.get("jd_requirement", ""))
        evidence = _esc(item.get("resume_evidence", ""))
        suggestion = _esc(item.get("suggestion", ""))
        
        # 状态文字颜色
        status_color = {"hit": "#16a34a", "weak": "#b45309", "gap": "#dc2626"}.get(status, "#666")
        status_label = {"hit": "✅ 覆盖", "weak": "⚠️ 较弱", "gap": "❌ 缺失"}.get(status, "?")
        status_icon = {"hit": "✅", "weak": "⚠️", "gap": "❌"}.get(status, "?")
        
        st.html(f"""
        <div class="card {card_class}">
          <!-- JD 要求 -->
          <div style="margin-bottom:12px">
            <div style="font-size:12px;color:#9ca3af;margin-bottom:4px">JD 第 {i} 条要求</div>
            <div style="font-size:16px;font-weight:700;color:#111827;line-height:1.6">{jd_req}</div>
          </div>

          <!-- 覆盖状态 -->
          <div style="display:flex;align-items:center;gap:10px;margin-bottom:12px">
            <span style="
              font-size:13px;font-weight:700;color:{status_color};
              background:{'#dcfce7' if status=='hit' else '#fef9c3' if status=='weak' else '#fee2e2'};
              padding:4px 12px;border-radius:20px">
              {status_icon}&nbsp;{status_label}
            </span>
            <span style="font-size:13px;color:#6b7280">
              {"简历已覆盖该要求，无需修改" if status=='hit' else "简历有相关经历，建议强化描述" if status=='weak' else "简历完全缺失该要求，建议补充"}
            </span>
          </div>

          {"<!-- 简历证据 -->" if evidence else ""}
          {f'''<div style="background:#f9fafb;border-radius:8px;padding:12px 14px;margin-bottom:12px">
            <div style="font-size:13px;font-weight:700;color:#374151;margin-bottom:4px">📌 简历中的对应内容</div>
            <div style="font-size:15px;color:#4b5563;line-height:1.6">{evidence}</div>
          </div>''' if evidence else ""}

          {"<!-- 改写建议 -->" if (suggestion and status != "hit") else ""}
          {f'''<div style="background:#eff6ff;border:1px solid #bfdbfe;border-radius:8px;padding:14px 16px">
            <div style="font-size:13px;font-weight:700;color:#1d4ed8;margin-bottom:6px">💡 改写建议（可直接采用或微调）</div>
            <div style="font-size:15px;color:#1e3a5f;line-height:1.8">{suggestion}</div>
          </div>''' if (suggestion and status != "hit") else ""}
        </div>""")

    st.divider()

    # ── 亮点 & 行动清单 ───────────────────────────────────
    col_a, col_b = st.columns(2, gap="large")

    with col_a:
        st.markdown("### 💪 你的优势亮点")
        strengths = result.get("strengths", [])
        for s in strengths:
            st.html(f"""<div class="card hit" style="padding:14px 18px;margin:6px 0;font-size:15px;color:#1f2937;line-height:1.6">✅ {_esc(s)}</div>""")

    with col_b:
        st.markdown("### ⚡ 优先行动清单")
        quick_wins = result.get("quick_wins", [])
        for j, qw in enumerate(quick_wins, 1):
            st.html(f"""<div class="card gap" style="padding:14px 18px;margin:6px 0;font-size:15px;color:#1f2937;line-height:1.6"><strong>#{j}</strong>&nbsp;&nbsp;{_esc(qw)}</div>""")

    st.divider()
    st.caption("🤖 由 DeepSeek AI 驱动 · JD Gap Analyzer MVP · 基于 SpeedyJob 改进方案")
